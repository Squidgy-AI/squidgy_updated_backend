"""
Background Text Processor
Handles downloading files from URLs and extracting text content
"""

import asyncio
import json
import logging
import tempfile
import os
import re
import base64
from typing import Optional, Dict, Any, List
import httpx
from pathlib import Path

# Text extraction libraries
try:
    import pdfplumber
    from io import BytesIO
    PyPDF2 = None  # Prefer pdfplumber over PyPDF2
except ImportError:
    pdfplumber = None
    try:
        import PyPDF2
        from io import BytesIO
    except ImportError:
        PyPDF2 = None

try:
    from docx import Document
except ImportError:
    Document = None

logger = logging.getLogger(__name__)

class TextExtractor:
    """Handles text extraction from different file types"""
    
    @staticmethod
    def extract_from_pdf(file_bytes: bytes) -> str:
        """
        Extract text from PDF bytes with proper reading order.
        Uses pdfplumber (preferred) for better layout preservation,
        falls back to PyPDF2 if pdfplumber not available.
        """
        # Try pdfplumber first (better reading order: top-left to bottom-right)
        if pdfplumber:
            try:
                pdf_file = BytesIO(file_bytes)
                text_content = []
                
                with pdfplumber.open(pdf_file) as pdf:
                    for i, page in enumerate(pdf.pages, 1):
                        # Extract words with positions for column-aware sorting
                        words = page.extract_words(
                            x_tolerance=3,
                            y_tolerance=3,
                            keep_blank_chars=False
                        )
                        
                        if words:
                            # Dynamic column detection based on word spacing analysis
                            # 1. Group words by line (same y position)
                            # 2. Calculate typical word spacing within each line
                            # 3. Detect column gaps as spacing significantly larger than typical
                            
                            # Sort words by y position first, then x
                            words_sorted = sorted(words, key=lambda w: (round(w['top']), w['x0']))
                            
                            # Group words into lines based on y position
                            lines_with_words = []
                            current_line = []
                            last_y = None
                            
                            for w in words_sorted:
                                y = round(w['top'])
                                if last_y is None or abs(y - last_y) <= 5:
                                    current_line.append(w)
                                else:
                                    if current_line:
                                        lines_with_words.append(current_line)
                                    current_line = [w]
                                last_y = y
                            if current_line:
                                lines_with_words.append(current_line)
                            
                            # Analyze spacing within lines to find typical word gap vs column gap
                            all_gaps = []
                            for line_words in lines_with_words:
                                if len(line_words) > 1:
                                    # Sort by x position within line
                                    line_words_sorted = sorted(line_words, key=lambda w: w['x0'])
                                    for k in range(len(line_words_sorted) - 1):
                                        gap = line_words_sorted[k + 1]['x0'] - line_words_sorted[k]['x1']
                                        if gap > 0:
                                            all_gaps.append(gap)
                            
                            # Determine if multi-column by finding outlier gaps
                            is_multi_column = False
                            column_boundary = None
                            
                            if all_gaps:
                                # Calculate median gap (typical word spacing)
                                sorted_gaps = sorted(all_gaps)
                                median_gap = sorted_gaps[len(sorted_gaps) // 2]
                                
                                # Column gap should be significantly larger than median (3x or more)
                                threshold = max(median_gap * 3, 20)  # At least 3x median or 20px
                                
                                # Find lines with large gaps (potential column separators)
                                column_gaps = []
                                for line_words in lines_with_words:
                                    if len(line_words) > 1:
                                        line_words_sorted = sorted(line_words, key=lambda w: w['x0'])
                                        for k in range(len(line_words_sorted) - 1):
                                            gap = line_words_sorted[k + 1]['x0'] - line_words_sorted[k]['x1']
                                            if gap > threshold:
                                                # Record the gap position (middle of the gap)
                                                gap_x = (line_words_sorted[k]['x1'] + line_words_sorted[k + 1]['x0']) / 2
                                                column_gaps.append(gap_x)
                                
                                # If multiple lines have large gaps at similar x positions, it's multi-column
                                if len(column_gaps) >= 2:
                                    # Find the most common gap position (cluster)
                                    column_gaps_sorted = sorted(column_gaps)
                                    # Use median of gap positions as column boundary
                                    column_boundary = column_gaps_sorted[len(column_gaps_sorted) // 2]
                                    is_multi_column = True
                            
                            if is_multi_column and column_boundary:
                                # Split words into left and right columns
                                left_words = [w for w in words if w['x1'] < column_boundary]
                                right_words = [w for w in words if w['x0'] > column_boundary]
                                
                                # Only proceed if both columns have content
                                if len(left_words) > 2 and len(right_words) > 2:
                                    # Sort each column by y position (top to bottom), then x
                                    left_words.sort(key=lambda w: (w['top'], w['x0']))
                                    right_words.sort(key=lambda w: (w['top'], w['x0']))
                                    
                                    # Build text from word list
                                    def words_to_text(word_list):
                                        if not word_list:
                                            return ""
                                        lines = []
                                        current_line = []
                                        last_top = word_list[0]['top']
                                        
                                        for w in word_list:
                                            if abs(w['top'] - last_top) > 5:
                                                if current_line:
                                                    lines.append(' '.join(current_line))
                                                current_line = [w['text']]
                                                last_top = w['top']
                                            else:
                                                current_line.append(w['text'])
                                        
                                        if current_line:
                                            lines.append(' '.join(current_line))
                                        
                                        return '\n'.join(lines)
                                    
                                    left_text = words_to_text(left_words)
                                    right_text = words_to_text(right_words)
                                    page_text = left_text + '\n\n' + right_text
                                else:
                                    page_text = page.extract_text(x_tolerance=3, y_tolerance=3)
                            else:
                                # Single column - use default extraction
                                page_text = page.extract_text(x_tolerance=3, y_tolerance=3)
                        else:
                            page_text = page.extract_text(x_tolerance=3, y_tolerance=3)
                        
                        if page_text and page_text.strip():
                            # Clean up whitespace
                            cleaned_text = re.sub(r'[ \t]+', ' ', page_text)
                            cleaned_text = re.sub(r'\n\s*\n+', '\n', cleaned_text)
                            cleaned_lines = [line.rstrip() for line in cleaned_text.split('\n')]
                            cleaned_text = '\n'.join(cleaned_lines).strip()
                            text_content.append(f"[Page {i}]\n{cleaned_text}")
                
                return '\n'.join(text_content).strip()
            except Exception as e:
                logger.error(f"pdfplumber extraction error: {e}")
                raise Exception(f"Failed to extract text from PDF: {str(e)}")
        
        # Fallback to PyPDF2
        if PyPDF2:
            try:
                pdf_file = BytesIO(file_bytes)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                text_content = []
                for i, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        cleaned = re.sub(r'\n\s*\n+', '\n', page_text.strip())
                        text_content.append(f"[Page {i}]\n{cleaned}")
                
                return '\n'.join(text_content).strip()
            except Exception as e:
                logger.error(f"PyPDF2 extraction error: {e}")
                raise Exception(f"Failed to extract text from PDF: {str(e)}")
        
        raise ImportError("No PDF library installed. Install pdfplumber or PyPDF2.")
    
    @staticmethod
    def extract_from_txt(file_bytes: bytes) -> str:
        """Extract text from TXT bytes with multiple encoding support"""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                return file_bytes.decode(encoding).strip()
            except UnicodeDecodeError:
                continue
        
        raise Exception("Failed to decode text file with supported encodings")
    
    @staticmethod
    def extract_from_json(file_bytes: bytes) -> str:
        """Extract text from JSON bytes - pretty-print the structured data"""
        encodings = ['utf-8', 'latin-1', 'cp1252']

        for encoding in encodings:
            try:
                raw = file_bytes.decode(encoding)
                parsed = json.loads(raw)
                return json.dumps(parsed, indent=2, ensure_ascii=False)
            except (UnicodeDecodeError, json.JSONDecodeError):
                continue

        raise Exception("Failed to parse JSON file")

    @staticmethod
    def extract_from_docx(file_bytes: bytes) -> str:
        """Extract text from DOCX bytes"""
        if not Document:
            raise ImportError("python-docx not installed")
        
        try:
            # Save bytes to temporary file for docx processing
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(file_bytes)
                temp_file_path = temp_file.name
            
            try:
                doc = Document(temp_file_path)
                text_content = []
                
                # Extract paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_content.append(paragraph.text.strip())
                
                # Extract table content
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text_content.append(' | '.join(row_text))
                
                return '\n'.join(text_content).strip()
            finally:
                # Clean up temporary file
                os.unlink(temp_file_path)
                
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")

    @staticmethod
    async def extract_from_image(file_bytes: bytes, file_name: str) -> str:
        """Extract text from image using OpenRouter vision model"""
        try:
            # Get OpenRouter API key from environment
            openrouter_api_key = os.getenv('OPENROUTER_API_KEY')
            if not openrouter_api_key:
                raise Exception("OPENROUTER_API_KEY not configured")

            # Convert image bytes to base64
            base64_image = base64.b64encode(file_bytes).decode('utf-8')

            # Detect image format from file extension
            file_ext = Path(file_name).suffix.lower()
            image_format_map = {
                '.png': 'image/png',
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg',
                '.gif': 'image/gif',
                '.bmp': 'image/bmp',
                '.webp': 'image/webp'
            }
            mime_type = image_format_map.get(file_ext, 'image/jpeg')

            # Prepare the API request
            async with httpx.AsyncClient(timeout=60.0) as client:
                response = await client.post(
                    "https://openrouter.ai/api/v1/chat/completions",
                    headers={
                        "Authorization": f"Bearer {openrouter_api_key}",
                        "Content-Type": "application/json"
                    },
                    json={
                        "model": "qwen/qwen2.5-vl-72b-instruct:free",
                        "messages": [
                            {
                                "role": "user",
                                "content": [
                                    {
                                        "type": "text",
                                        "text": "Extract all text from this image. Return only the extracted text without any additional commentary or formatting. If there is no text in the image, return 'No text found in image'."
                                    },
                                    {
                                        "type": "image_url",
                                        "image_url": {
                                            "url": f"data:{mime_type};base64,{base64_image}"
                                        }
                                    }
                                ]
                            }
                        ]
                    }
                )

                response.raise_for_status()
                result = response.json()

                # Extract the text from the response
                extracted_text = result.get('choices', [{}])[0].get('message', {}).get('content', '')

                if not extracted_text or extracted_text.strip() == '':
                    return f"[Image file: {file_name}] - No text extracted"

                return extracted_text.strip()

        except Exception as e:
            logger.error(f"Image text extraction error: {e}")
            # Fallback message if extraction fails
            return f"[Image file: {file_name}] - Text extraction failed: {str(e)}"


class BackgroundTextProcessor:
    """Handles background processing of file text extraction"""

    def __init__(self, supabase_client):
        self.supabase = supabase_client
        self.text_extractor = TextExtractor()
    
    async def download_file(self, file_url: str) -> bytes:
        """Download file from URL and return bytes"""
        try:
            async with httpx.AsyncClient(timeout=60.0, follow_redirects=True) as client:
                response = await client.get(file_url)
                response.raise_for_status()
                return response.content
        except httpx.TimeoutException:
            raise Exception("File download timeout")
        except httpx.HTTPStatusError as e:
            raise Exception(f"Failed to download file: HTTP {e.response.status_code}")
        except Exception as e:
            raise Exception(f"File download error: {str(e)}")
    
    @staticmethod
    def chunk_text(text: str, chunk_size: int = 4000, chunk_overlap: int = 400) -> List[str]:
        """Split text into chunks with overlap, breaking at sentence/paragraph boundaries."""
        if not text or not text.strip():
            return []

        # If text fits in one chunk, return as-is
        if len(text) <= chunk_size:
            return [text.strip()]

        chunks = []
        start = 0

        while start < len(text):
            end = start + chunk_size

            # If this is the last chunk, take the rest
            if end >= len(text):
                chunk = text[start:].strip()
                if chunk:
                    chunks.append(chunk)
                break

            # Try to break at paragraph boundary (double newline)
            segment = text[start:end]
            break_pos = segment.rfind('\n\n')

            # Fallback: break at single newline
            if break_pos == -1 or break_pos < chunk_size * 0.3:
                break_pos = segment.rfind('\n')

            # Fallback: break at sentence end (. ! ?)
            if break_pos == -1 or break_pos < chunk_size * 0.3:
                sentence_match = None
                for m in re.finditer(r'[.!?]\s', segment):
                    if m.start() >= chunk_size * 0.3:
                        sentence_match = m
                if sentence_match:
                    break_pos = sentence_match.end()

            # Fallback: break at space
            if break_pos == -1 or break_pos < chunk_size * 0.3:
                break_pos = segment.rfind(' ')

            # Last resort: hard cut
            if break_pos == -1 or break_pos < chunk_size * 0.3:
                break_pos = chunk_size

            chunk = text[start:start + break_pos].strip()
            if chunk:
                chunks.append(chunk)

            # Move start forward, accounting for overlap
            start = start + break_pos - chunk_overlap
            if start < 0:
                start = 0

        return chunks

    async def extract_text(self, file_bytes: bytes, file_name: str) -> str:
        """Extract text based on file extension"""
        file_ext = Path(file_name).suffix.lower()

        # Handle image files - extract text using OpenRouter vision model
        if file_ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp']:
            return await self.text_extractor.extract_from_image(file_bytes, file_name)

        if file_ext == '.pdf':
            return self.text_extractor.extract_from_pdf(file_bytes)
        elif file_ext in ['.txt', '.md', '.markdown']:
            return self.text_extractor.extract_from_txt(file_bytes)
        elif file_ext in ['.json']:
            return self.text_extractor.extract_from_json(file_bytes)
        elif file_ext in ['.docx', '.doc']:
            return self.text_extractor.extract_from_docx(file_bytes)
        else:
            # Fallback: try to read as plain text
            try:
                return self.text_extractor.extract_from_txt(file_bytes)
            except Exception:
                raise Exception(f"Unsupported file type: {file_ext}")
    
    async def update_neon_record_ids(self, file_id: str, neon_record_ids: list):
        """Update the neon_record_ids in Supabase after saving to Neon"""
        try:
            # Convert to strings to avoid scientific notation for large integers in JSON
            # Neon IDs are bigint/serial which can be very large
            serializable_ids = [str(id) for id in neon_record_ids]
            
            result = self.supabase.table("firm_users_knowledge_base").update({
                "neon_record_ids": serializable_ids,
                "updated_at": "now()"
            }).eq("file_id", file_id).execute()
            
            if not result.data:
                logger.error(f"Failed to update neon_record_ids for file_id: {file_id}")
            else:
                logger.info(f"Updated neon_record_ids for {file_id}: {len(serializable_ids)} records")
                
        except Exception as e:
            logger.error(f"Database update error for {file_id}: {e}")
    
    async def generate_embedding(self, text: str) -> Optional[str]:
        """
        Generate embedding using OpenRouter API.
        Returns formatted string for PostgreSQL vector: [0.1,0.2,...]
        
        Note: Uses same format as generate_embedding_for_kb() in main.py
        """
        try:
            OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
            if not OPENROUTER_API_KEY:
                logger.error("OPENROUTER_API_KEY not configured")
                return None
            
            async with httpx.AsyncClient(timeout=60.0) as client:
                response = await client.post(
                    "https://openrouter.ai/api/v1/embeddings",
                    headers={
                        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                        "Content-Type": "application/json"
                    },
                    json={
                        "model": "openai/text-embedding-3-small",
                        "input": text[:8000]
                    }
                )
                
                if response.status_code == 200:
                    data = response.json()
                    embedding = data.get("data", [{}])[0].get("embedding", [])
                    if embedding and len(embedding) > 0:
                        vector_str = '[' + ','.join(str(x) for x in embedding) + ']'
                        return vector_str
                else:
                    logger.error(f"Embedding API error: {response.status_code}")
                    
        except Exception as e:
            logger.error(f"Error generating embedding: {str(e)}")
        
        return None
    
    async def save_chunks_to_neon(
        self,
        user_id: str,
        agent_id: str,
        file_name: str,
        file_url: str,
        extracted_text: str,
        source: str = "chat",
        chunk_size: int = 4000,
        chunk_overlap: int = 400
    ) -> List[int]:
        """
        Save extracted text chunks to Neon DB with embeddings.
        Returns list of Neon record IDs.
        """
        import asyncpg
        from datetime import datetime
        
        NEON_DB_HOST = os.getenv("NEON_DB_HOST")
        NEON_DB_USER = os.getenv("NEON_DB_USER")
        NEON_DB_PASSWORD = os.getenv("NEON_DB_PASSWORD")
        NEON_DB_NAME = os.getenv("NEON_DB_NAME", "neondb")
        
        if not all([NEON_DB_HOST, NEON_DB_USER, NEON_DB_PASSWORD]):
            logger.error("Neon DB credentials not configured")
            return []
        
        # Chunk the text
        chunks = self.chunk_text(extracted_text, chunk_size, chunk_overlap)
        if not chunks:
            logger.warning(f"No chunks generated for {file_name}")
            return []
        
        logger.info(f"Generated {len(chunks)} chunks for {file_name}")
        
        neon_record_ids = []
        
        try:
            conn = await asyncpg.connect(
                host=NEON_DB_HOST,
                user=NEON_DB_USER,
                password=NEON_DB_PASSWORD,
                database=NEON_DB_NAME,
                ssl="require"
            )
            
            try:
                created_at = datetime.utcnow()
                
                for i, chunk in enumerate(chunks, start=1):
                    # Format document with metadata
                    formatted_doc = f"[Source: {file_name} | Chunk {i}/{len(chunks)}]\n\n{chunk}"
                    
                    # Generate embedding
                    embedding = await self.generate_embedding(chunk)
                    
                    if embedding:
                        # Insert into Neon DB
                        insert_query = """
                            INSERT INTO user_vector_knowledge_base
                            (user_id, agent_id, document, embedding, category, source, file_name, file_url, created_at, updated_at)
                            VALUES ($1, $2, $3, $4::vector, $5, $6, $7, $8, $9, $10)
                            RETURNING id
                        """
                        
                        result = await conn.fetchrow(
                            insert_query,
                            user_id,
                            agent_id,
                            formatted_doc,
                            embedding,
                            'documents',
                            source,
                            file_name,
                            file_url,
                            created_at,
                            created_at
                        )
                        
                        if result:
                            neon_record_ids.append(result['id'])
                            logger.debug(f"Saved chunk {i}/{len(chunks)} to Neon (id: {result['id']})")
                    else:
                        logger.warning(f"Failed to generate embedding for chunk {i}/{len(chunks)}")
                
                logger.info(f"Saved {len(neon_record_ids)} chunks to Neon for {file_name}")
                
            finally:
                await conn.close()
                
        except Exception as e:
            logger.error(f"Error saving to Neon: {str(e)}")
        
        return neon_record_ids
    
    async def process_file(self, file_id: str):
        """Main processing function for a file - extracts text and saves to Neon"""
        try:
            # Get file record from Supabase firm_users_knowledge_base
            result = self.supabase.table("firm_users_knowledge_base").select(
                "file_id, file_name, file_url, firm_user_id, agent_id, source, neon_record_ids"
            ).eq("file_id", file_id).execute()

            if not result.data:
                logger.error(f"File record not found: {file_id}")
                return

            file_record = result.data[0]
            
            # Skip if already has Neon records (already processed)
            existing_neon_ids = file_record.get("neon_record_ids", [])
            if existing_neon_ids and len(existing_neon_ids) > 0:
                logger.info(f"File {file_id} already has {len(existing_neon_ids)} Neon records, skipping")
                return

            # Download file
            logger.info(f"Downloading file: {file_record['file_url']}")
            file_bytes = await self.download_file(file_record["file_url"])

            # Extract text
            logger.info(f"Extracting text from: {file_record['file_name']}")
            extracted_text = await self.extract_text(file_bytes, file_record["file_name"])

            if not extracted_text.strip():
                logger.error(f"No text content found in file {file_id}")
                return

            # Save chunks to Neon and get record IDs
            neon_record_ids = await self.save_chunks_to_neon(
                user_id=file_record["firm_user_id"],
                agent_id=file_record["agent_id"],
                file_name=file_record["file_name"],
                file_url=file_record["file_url"],
                extracted_text=extracted_text,
                source=file_record.get("source", "chat")
            )

            # Update Supabase with Neon record IDs
            if neon_record_ids:
                await self.update_neon_record_ids(file_id, neon_record_ids)
                logger.info(f"Successfully processed file {file_id} - {len(neon_record_ids)} chunks saved")
            else:
                logger.error(f"No chunks saved to Neon for file {file_id}")

        except Exception as e:
            logger.error(f"Processing failed for {file_id}: {str(e)}")


# Global processor instance
_background_processor = None

def get_background_processor():
    """Get the global background processor instance"""
    global _background_processor
    return _background_processor

def initialize_background_processor(supabase_client):
    """Initialize the global background processor"""
    global _background_processor
    _background_processor = BackgroundTextProcessor(supabase_client)
    return _background_processor